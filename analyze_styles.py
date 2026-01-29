from docx import Document

def analyze_paragraph_styles(docx_path):
    """分析文档中段落的样式和字号"""
    print(f"正在分析文档: {docx_path}")
    print("=" * 60)
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"读取文档失败: {e}")
        return
    
    style_count = {}
    font_size_count = {}
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        text = paragraph.text.strip()[:30]
        
        style_name = ""
        if paragraph.style:
            style_name = paragraph.style.name
        
        if style_name not in style_count:
            style_count[style_name] = 0
        style_count[style_name] += 1
        
        for run in paragraph.runs:
            if run.font.size:
                font_size = run.font.size.pt
                if font_size not in font_size_count:
                    font_size_count[font_size] = 0
                font_size_count[font_size] += 1
                break
            elif run.style and run.style.font and run.style.font.size:
                font_size = run.style.font.size.pt
                if font_size not in font_size_count:
                    font_size_count[font_size] = 0
                font_size_count[font_size] += 1
                break
    
    print("\n样式统计:")
    print("-" * 60)
    for style, count in sorted(style_count.items(), key=lambda x: x[1], reverse=True):
        print(f"  {style}: {count} 个段落")
    
    print("\n字号统计:")
    print("-" * 60)
    for size, count in sorted(font_size_count.items(), key=lambda x: x[0]):
        print(f"  {size}pt: {count} 个段落")
    
    print("\n前20个段落详情:")
    print("-" * 60)
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        text = paragraph.text.strip()[:50]
        
        style_name = ""
        if paragraph.style:
            style_name = paragraph.style.name
        
        font_size = "默认"
        for run in paragraph.runs:
            if run.font.size:
                font_size = f"{run.font.size.pt}pt"
                break
            elif run.style and run.style.font and run.style.font.size:
                font_size = f"{run.style.font.size.pt}pt"
                break
        
        print(f"\n段落 {para_idx + 1}:")
        print(f"  样式: {style_name}")
        print(f"  字号: {font_size}")
        print(f"  内容: {text}")
        
        if para_idx >= 19:
            break

if __name__ == '__main__':
    analyze_paragraph_styles("水军话术(2).docx")

